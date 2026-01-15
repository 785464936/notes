#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
创建符合格式要求的课程设计说明书
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建新文档
doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 添加封面页
def add_cover_page(doc):
    # 校徽占位符
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[校徽]")
    run.font.size = Pt(12)

    # 校名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("产品制造综合实践")
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 设计题目
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("设计题目：CA6140车床后托架工艺规程及夹具设计")
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True

    # 空行
    for _ in range(5):
        doc.add_paragraph()

    # 作者信息
    def add_centered_text(text, font_name, font_size, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.bold = bold
        set_paragraph_spacing(p, before=Pt(6), after=Pt(6), line_spacing=Pt(20))

    add_centered_text("姓　　名：时宝鑫", '仿宋', 16)
    add_centered_text("学　　号：", 'Times New Roman', 16)
    add_centered_text("学　　院：机械工程学院", '仿宋', 16)
    add_centered_text("专　　业：增材制造工程", '仿宋', 16)
    add_centered_text("指导教师：齐习娟", '仿宋', 16)

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 日期
    add_centered_text("二〇二六年一月", '宋体', 16)

    # 分页
    doc.add_page_break()

def set_paragraph_spacing(paragraph, before=None, after=None, line_spacing=None):
    """设置段落间距"""
    pPr = paragraph._element.pPr
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)

    if before is not None or after is not None:
        spacing = OxmlElement('w:spacing')
        if before is not None:
            spacing.set(qn('w:before'), str(before))
        if after is not None:
            spacing.set(qn('w:after'), str(after))
        pPr.append(spacing)

    if line_spacing is not None:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(line_spacing))
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)

def add_heading(doc, text, level):
    """添加标题"""
    if level == 1:
        p = doc.add_paragraph(text, style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=Pt(40), after=Pt(20), line_spacing=Pt(20))
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(15)
            run.bold = True
    elif level == 2:
        p = doc.add_paragraph(text, style='Heading 2')
        set_paragraph_spacing(p, before=Pt(24), after=Pt(6), line_spacing=Pt(20))
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
            run.bold = True
    elif level == 3:
        p = doc.add_paragraph(text, style='Heading 3')
        set_paragraph_spacing(p, before=Pt(12), after=Pt(6), line_spacing=Pt(20))
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(13)
            run.bold = True

def add_paragraph_text(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符缩进
    set_paragraph_spacing(p, before=Pt(0), after=Pt(0), line_spacing=Pt(20))
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

# 创建封面
add_cover_page(doc)

# 添加目录页
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("目　　录")
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(15)
run.bold = True
set_paragraph_spacing(p, before=Pt(40), after=Pt(20), line_spacing=Pt(20))

# 目录内容（占位符，实际使用时需要根据内容生成）
toc_items = [
    ("1　引言", 1),
    ("　　1.1　设计背景与意义", 1),
    ("2　零件的工艺性分析", 1),
    ("　　2.1　零件的作用及技术要求分析", 1),
    ("　　2.2　审查零件结构工艺性分析", 1),
    ("　　2.3　确定毛坯类型和形状", 1),
    ("3　零件的加工工艺设计", 1),
    ("　　3.1　定位基面的选择", 1),
    ("　　3.2　工序合理组成", 1),
    ("　　3.3　加工工艺路线", 1),
    ("　　3.4　工艺方案的比较与分析", 1),
    ("　　3.5　刀具的选择", 1),
    ("　　3.6　切削用量的计算", 1),
    ("4　粗镗孔夹具设计", 1),
    ("　　4.1　设计要求", 1),
    ("　　4.2　夹具的基准选择与计算", 1),
    ("　　4.3　定位误差的分析", 1),
    ("结论", 1),
    ("参考文献", 1),
    ("致谢", 1),
]

for item, level in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74 * (level - 1))
    set_paragraph_spacing(p, before=Pt(6), after=Pt(0), line_spacing=Pt(20))
    run = p.add_run(item)
    if level == 1:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    else:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

doc.add_page_break()

# 添加正文内容
add_heading(doc, "1　引言", 1)
add_paragraph_text(doc, "机械制造工艺学课程设计是我们在学完大学的大部分课程后进行的，是我们对大学三年的学习的一次深入的综合性的总考核，也是一次理论联系实际的训练，这次设计使我们能综合运用机械制造工艺学中的基本理论，并结合实习中学到的实践知识，独立地分析和解决工艺问题，初步具备了设计一个中等复杂程度零件（CA6140车床后托架）的工艺规程的能力和运用夹具设计的基本原理和方法，拟订夹具设计方案，完成夹具结构设计的能力，也是熟悉和运用有关手册、图表等技术资料及编写技术文件等基本技能的一次实践机会。")
add_paragraph_text(doc, "因此，它在我们大学生活中占有重要地位。就我个人而言，我也希望通过这次设计对自己未来将从事的工作进行一次适应性心理，从中锻炼自己分析问题，解决问题的能力，对未来的工作发展打下一个良好的基础。")
add_paragraph_text(doc, "本次课程设计，主要围绕小组任务——CA6140车床后托架设计进行任务的分配，涉及到毛坯件成型，毛坯件铸造方法及基本雏形的拟定，加工余量的确定，工序分配，夹具设计，数据计算等一系列系统的设计思路。全组共7人，分工明确，合作配合默契。")
add_paragraph_text(doc, "在设计过程中，对于一些元件的选取，工艺方案的拟定，工序流程的安排，小组讨论中都存在一些显著的问题，基础知识不扎实，讨论效果不理想，思路不清晰，等一些因素导致设计进展不理想，但在指导老师的指导，纠错和引领下，相对满意的完成了本次课程设计的训练任务。由于能力所限，设计尚有许多不足之处，恳请老师们给予指教。望在细节处多批评，万分感谢。")

add_heading(doc, "2　零件的工艺性分析", 1)
add_heading(doc, "2.1　零件的作用及技术要求分析", 2)
add_paragraph_text(doc, "课程设计题目所给定的零件是CA6140车床后托架，CA6140机床后托架是CA6140机床的一个重要零件，因为其零件尺寸较小，结构形状也不是很复杂，但侧面三杠孔和底面的精度要求较高，此外还有顶面的四孔要求加工，但是对精度要求不是很高。后托架上的侧面三杠孔的粗糙度要求都是Ra1.6，所以都要求精加工。其三杠孔的中心线和底平面有平面度的公差要求等。因为其尺寸精度、几何形状精度和相互位置精度，以及各表面的表面质量均影响机器或部件的装配质量，进而影响其性能与工作寿命，因此它的加工是非常关键和重要的。")
add_paragraph_text(doc, "后托架主要起支撑作用。在加工长轴类零件时，后托架能够对零件的后端提供稳定的支撑。因为长轴在加工过程中容易出现下垂，后托架可以防止零件因自身重力而产生变形，确保零件的中心轴线与车床主轴轴线尽可能重合。")
add_paragraph_text(doc, "此外，后托架还能辅助提高加工精度。当刀具对工件进行切削加工时，会产生切削力。后托架能够帮助工件抵抗切削力，减少工件的振动和偏移，从而提升零件的圆柱度、同轴度等加工精度指标，保证加工质量。")

add_heading(doc, "2.2　审查零件结构工艺性分析", 2)
add_heading(doc, "2.2.1　结构特点", 3)
add_paragraph_text(doc, "形状复杂：整体呈不规则形状，有多个平面、孔系和凸台。例如，有用于安装的底面和侧面，以及支撑轴类零件的半圆形凸台等。")
add_paragraph_text(doc, "孔系分布：包含不同直径和深度的孔，如光孔、螺纹孔等，且各孔的位置精度要求较高，需保证相互之间的平行度、垂直度和同轴度，以满足与其他部件的装配要求。")
add_paragraph_text(doc, "壁薄不均：部分区域壁薄，在加工过程中易变形，如一些连接部位和凸台边缘，对加工工艺提出挑战。")

add_heading(doc, "2.2.2　工艺性分析", 3)
add_paragraph_text(doc, "加工表面：底面和侧面要求较高的平面度和表面粗糙度，以保证安装精度；孔的加工需保证尺寸精度和形状精度，否则影响后托架的使用性能。")
add_paragraph_text(doc, "定位基准：选择合适的定位基准至关重要。通常以底面和侧面作为粗基准，后续加工中以已加工表面为精基准，保证各加工表面之间的位置精度关系，减少累积误差。")
add_paragraph_text(doc, "加工顺序：遵循先面后孔、先粗后精原则。先加工平面，为孔加工提供稳定可靠的基准；粗加工去除余量，再进行半精和精加工，逐步提高精度，降低加工难度和成本，保证加工质量和效率。")

add_heading(doc, "2.3　确定毛坯类型和形状", 2)
add_paragraph_text(doc, "毛坯的种类一般根据零件的设计要求、结构形状、使用要求和生产类型等进行选择。形状复杂的零件毛坯，如箱体、机架底座、床身、壳体等，一般采用铸件。机械性能要求高、形状较为简单、有一定批量的重要零件毛坯，如曲轴、连杆等，一般选择锻件。结构简单的零件毛坯，如轴类、套类、盘类、板类、长条类零件，一般采用型材。")
add_paragraph_text(doc, "查机械制造技术基础课程设计表2.6得毛坯制造方式有很多，但考虑到本次课程设计题目的要求以及具体的应用情况，毛坯制造采用铸造方式，根据毛坯的尺寸大小以及产量，效率，应用场合等因素，选用金属铸造，对毛坯件进行成型铸造。")
add_paragraph_text(doc, "后托架主要起支撑作用，要承受一定的力，需要有足够的强度和刚度。灰铸铁（HT200）这类材料制成的铸件毛坯比较合适，其抗压强度和减震性良好，能满足后托架的使用要求。")
add_paragraph_text(doc, "从形状上看，为便于加工，毛坯形状要尽量接近成品形状。CA6140后托架结构较为复杂，有多个安装平面、孔系以及支撑部分。采用铸件可以铸出比较复杂的形状，像内部的加强肋、各种凸台等部分都能一体成型，减少后续加工量。")

add_heading(doc, "3　零件的加工工艺设计", 1)
add_heading(doc, "3.1　定位基面的选择", 2)
add_paragraph_text(doc, "基面的选择是工艺规程设计中的重要工作之一，基面选择的正确与合理，可以使加工质量得到保证，生产率得以提高。否则，加工工艺过程中会问题百出，更有甚者，还会造成零件大批报废，使生产无法正常进行。")
add_paragraph_text(doc, "工序I是铣车床后托架底面，而侧面通常是在铸造或者粗加工时比较容易保证精度的平面。以它为基准能够更好地保证底面与侧面的垂直度。因为在加工过程中，垂直度是一个很重要的形位公差要求，这对于后托架后续的安装以及和其他部件的配合至关重要。")
add_paragraph_text(doc, "其次，从装夹角度考虑，侧面定位方便、稳定。可以使用简单的夹具（如平口虎钳）进行装夹，并且在装夹过程中，由于侧面面积相对较大，能够提供足够的定位面积，使得工件在铣削底面时不易产生位移和振动，有利于提高铣削质量，保证底面的平整度。")

add_heading(doc, "3.2　工序合理组成", 2)
add_paragraph_text(doc, "由于底面与Φ40孔有平行度要求，即Φ40孔加工时需以底面作为基准，以满足与地面的平行度要求，而且相对于后面的加工工序，用底面作为后续工序的加工基准更加合理牢固，减少了加工误差，所以需要将铣底面作为第一道工序。")
add_paragraph_text(doc, "而最后攻丝，是为了防止加工过程中需要Φ6孔的定位，如果先进行攻丝，则被攻丝面无法作为基准面，否则无法满足加工精度要求，也会破坏螺纹线。")

add_heading(doc, "3.3　加工工艺路线", 2)
add_paragraph_text(doc, "查《机械制造技术基础课程设计》表2.25平面加工方案的经济精度和表面粗糙度，由于金属铸造公差等级IT为8-10，所以加工方案选择精铣和半精铣结合。")
add_paragraph_text(doc, "从加工精度方面来看，粗铣可以快速去除大量毛坯余量。后托架毛坯通常余量较多，粗铣能高效地将大部分余量去除，使工件接近最终尺寸。半精铣可以进一步提高加工精度。粗铣后的表面粗糙度较大，尺寸精度也较低，半精铣能够减小表面粗糙度，提高底面的平面度，更好地满足设计要求。")
add_paragraph_text(doc, "从加工效率方面考虑，粗铣采用较大的切削深度和进给量，能够快速去除余量，节省加工时间。如果只用精铣的方式来加工，由于切削深度小，要去除大量余量会花费很长时间。半精铣切削用量适中，在保证加工精度的同时，也不会过于降低加工效率。")
add_paragraph_text(doc, "从刀具寿命角度分析，粗铣时，由于切削深度和进给量较大，刀具磨损相对较快，但因为主要是去除余量，对表面质量要求不高，所以一定程度的刀具磨损是可以接受的。半精铣时，切削用量较小，刀具磨损相对缓慢，有利于获得良好的加工表面，并且可以延长刀具的使用寿命。")

add_heading(doc, "3.4　工艺方案的比较与分析", 2)
add_paragraph_text(doc, "加工精度方面，粗铣和半精铣，粗铣主要是快速去除毛坯余量，半精铣在粗铣基础上进一步提高精度。它能使工件的尺寸精度达到IT8-IT11，表面粗糙度Ra值达到3.2-12.5μm左右。对于CA6140车床后托架底面，这种精度可以满足后续加工或装配的基本要求。")
add_paragraph_text(doc, "加工效率方面，粗铣和半精铣，粗铣采用较大的切削用量，能快速去除余量。半精铣的切削用量比精铣稍大，加工时间相对较短。整个粗铣和半精铣过程在保证一定精度的同时，效率较高。")
add_paragraph_text(doc, "刀具磨损和成本方面，粗铣和半精铣，粗铣时刀具磨损快，但由于半精铣的切削条件相对精铣没那么苛刻，对刀具的精度要求稍低，刀具成本相对较低。而且半精铣可以在一定程度上修正粗铣时刀具磨损造成的表面缺陷。")

add_heading(doc, "3.5　刀具的选择", 2)
add_paragraph_text(doc, "由于加工底面有一定的精度要求，且毛坯尺寸不足3mm，由于受被吃刀量的限制，查书选择了高速钢圆柱铣刀。高速钢圆柱铣刀的硬度较高，能有效切削CA6140车床后托架。它可以承受一定的切削力，在铣削过程中保持刀刃的锋利度，对于加工后托架底面、侧面等部位的余量去除比较有效。")
add_paragraph_text(doc, "加工精度较高，高速钢圆柱铣刀的制造精度较高，可以加工出尺寸精度较高的表面。对于CA6140后托架来说，在铣削安装平面、支撑面等部位时，能够保证较好的平面度和表面粗糙度。一般可以达到表面粗糙度Ra3.2-12.5μm左右，满足后托架在车床中正常使用的精度要求。")
add_paragraph_text(doc, "刀具韧性好，高速钢圆柱铣刀的韧性优于一些硬质合金刀具。在铣削CA6140车床后托架时，由于后托架毛坯可能存在一些硬点或者加工过程中出现轻微的振动等情况，高速钢立铣刀不容易发生崩刃现象。")
add_paragraph_text(doc, "刃磨方便，当高速钢圆柱铣刀的刀刃磨损后，可以比较方便地进行刃磨。这对于加工CA6140车床后托架这种可能需要多次铣削操作的工件来说很重要。通过刃磨，可以恢复刀具的切削性能，延长刀具的使用寿命，降低加工成本。")

add_heading(doc, "3.6　切削用量的计算", 2)
add_heading(doc, "3.6.1　加工余量的确定", 3)
add_paragraph_text(doc, "零件材料为HT200，硬度170～220HBS，毛坯重量为3.05kg。生产类型为中批及以上生产，采用砂型铸造毛坯。")
add_paragraph_text(doc, "根据上述原始资料及加工工艺，分别确定各加工表面的机械加工余量、工序尺寸及毛坯尺寸如下：")
add_paragraph_text(doc, "（1）铣底面A为120mm方向的加工余量及公差")
add_paragraph_text(doc, "查《机械加工工艺手册》（以下简称《工艺手册》）表3.1-24，成批和大量生产铸件的尺寸公差等级，查得铸件尺寸公差等级CT分为8-10级，选用8级。")
add_paragraph_text(doc, "查表3.1-21，得铸件尺寸公差为1.8mm。")
add_paragraph_text(doc, "查表3.1-27，得此铸件机械加工余量等级为D-F级，选用F级。")
add_paragraph_text(doc, "查表3.1-26，得加工余量值为2mm。")
add_paragraph_text(doc, "（2）底面毛坯尺寸计算")
add_paragraph_text(doc, "底边毛坯厚度=0.5公差+加工余量=2.9mm")

add_heading(doc, "3.6.2　粗镗Φ25孔的切削用量计算", 3)
add_paragraph_text(doc, "工序VII 粗镗Φ25+0.3 0孔")
add_paragraph_text(doc, "切削深度ap=0.97mm，孔径d0=23.35mm")
add_paragraph_text(doc, "进给量f根据《机械加工工艺手册》表1.2-33，刀杆伸出长度取200mm，切削深度为0.97mm。因此确定进给量f=0.6mm/r")
add_paragraph_text(doc, "切削速度v参照《机械加工工艺手册》表1.2-33，取v=60m/min")
add_paragraph_text(doc, "机床主轴转速n，由式（1.2）有：n=1000V/πd0=1000×60/3.14×23.35≈817.93r/min，取n=1000r/min")
add_paragraph_text(doc, "实际切削速度V'，由式（1.2）有：v'=πd0n/1000=3.14×23.35×1000/1000×60≈1.22m/s")
add_paragraph_text(doc, "工作台每分钟进给量fm，由式（1.7）有：fm=f×n=0.6×1000=600mm/min")
add_paragraph_text(doc, "被切削层长度l=60mm")
add_paragraph_text(doc, "刀具切入长度l1，由式（1.6）有：l1=ap/tankr+(2~3)=1.7/tan30°+2≈3.69mm")
add_paragraph_text(doc, "刀具切出长度l2=3~5mm，取l2=4mm，行程次数i=1")
add_paragraph_text(doc, "机动时间tj1=(l+l1+l2)/fm=(60+3.69+4)/(600×1)≈0.11min")

add_heading(doc, "4　粗镗孔夹具设计", 1)
add_heading(doc, "4.1　设计要求", 2)
add_paragraph_text(doc, "本夹具主要用来粗镗Φ25.5mm孔。加工时除了要满足粗糙度要求外，还应满足孔轴线对底平面的平行度公差要求。为了保证技术要求，最关键是找到定位基准。同时，应考虑如何提高劳动生产率和降低劳动强度。")

add_heading(doc, "4.2　夹具的基准选择与计算", 2)
add_heading(doc, "4.2.1　定位基准的选择", 3)
add_paragraph_text(doc, "由零件图可知：孔的轴线与底平面有平行度公差要求，在对孔进行加工前，底平面进行了精铣加工。因此，选底平面为定位精基准来满足平行度公差要求。")
add_paragraph_text(doc, "同时孔的轴线间有位置公差，选择左侧面为定位基准来设计镗模，从而满足孔轴线间的位置公差要求。工件定位用底平面和两个侧面来限制六个自由度。")

add_heading(doc, "4.3　定位误差的分析", 2)
add_paragraph_text(doc, "（此部分内容待补充）")

add_heading(doc, "结论", 1)
add_paragraph_text(doc, "通过本次课程设计，我系统地学习了机械制造工艺规程设计的基本方法和步骤，掌握了夹具设计的基本原理。在设计过程中，我深入理解了CA6140车床后托架的加工工艺特点，合理选择了加工方法、定位基准和工艺参数，完成了工艺规程和夹具方案的设计。")
add_paragraph_text(doc, "通过本次设计，我不仅巩固了所学理论知识，还提高了分析问题和解决问题的能力，为今后的工作和学习打下了良好的基础。")

add_heading(doc, "参考文献", 1)
add_paragraph_text(doc, "[1] 机械加工工艺手册[M].北京:机械工业出版社,1998.")
add_paragraph_text(doc, "[2] 机械制造技术基础课程设计[M].北京:机械工业出版社,2010.")
add_paragraph_text(doc, "[3] 机床夹具设计手册[M].北京:机械工业出版社,2009.")

add_heading(doc, "致谢", 1)
add_paragraph_text(doc, "本次课程设计是在指导老师齐习娟的悉心指导下完成的。从选题、方案确定到最终完成，齐老师都给予了耐心的指导和帮助，使我能够顺利完成设计任务。在此，我向齐老师表示衷心的感谢！")
add_paragraph_text(doc, "同时，感谢小组其他成员的协作和帮助，大家在讨论中相互学习、共同进步，使我受益匪浅。")
add_paragraph_text(doc, "最后，感谢所有关心和帮助过我的老师和同学！")

# 保存文档
output_path = "D:\\NOTES\\01Project\\课设-机械产品设计\\课程设计说明书_格式化.docx"
doc.save(output_path)
print(f"文档已保存至: {output_path}")